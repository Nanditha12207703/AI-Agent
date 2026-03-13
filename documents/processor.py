"""
documents/processor.py - Fixed: pandas optional, uses csv module as fallback
"""
import io
import os
import csv
from pathlib import Path
from typing import Tuple
from loguru import logger

# pandas is optional - fall back to csv module if not available
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logger.warning("pandas not available - CSV/XLSX will use basic extraction")


class DocumentProcessor:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".csv", ".xlsx", ".txt", ".md"}

    def extract_text(self, file_path: str, filename: str) -> Tuple[str, str]:
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{ext}'. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}")
        if ext == ".pdf":
            return self._extract_pdf(file_path), "pdf"
        elif ext in (".docx", ".doc"):
            return self._extract_docx(file_path), "docx"
        elif ext == ".csv":
            return self._extract_csv(file_path), "csv"
        elif ext == ".xlsx":
            return self._extract_xlsx(file_path), "xlsx"
        else:
            return self._extract_text(file_path), "txt"

    def _extract_pdf(self, file_path: str) -> str:
        try:
            import pypdf
            text_parts = []
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {i+1}]\n{page_text}")
            full_text = "\n\n".join(text_parts)
            logger.info(f"PDF extracted: {len(full_text)} chars")
            return full_text or "[PDF appears to be empty or image-only]"
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Could not extract PDF content: {e}")

    def _extract_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            table_texts = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    table_texts.append("\n".join(rows))
            full = "\n\n".join(paragraphs)
            if table_texts:
                full += "\n\n[Tables]\n" + "\n\n".join(table_texts)
            logger.info(f"DOCX extracted: {len(full)} chars")
            return full or "[Document appears to be empty]"
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Could not extract DOCX content: {e}")

    def _extract_csv(self, file_path: str) -> str:
        try:
            if HAS_PANDAS:
                df = pd.read_csv(file_path, nrows=500)
                summary = f"CSV File Summary\nRows: {len(df)} | Columns: {list(df.columns)}\n\n"
                summary += "Sample Data (first 20 rows):\n"
                summary += df.head(20).to_string(index=False)
                numeric_cols = df.select_dtypes(include="number").columns.tolist()
                if numeric_cols:
                    summary += "\n\nNumeric Summary:\n"
                    summary += df[numeric_cols].describe().to_string()
                return summary
            else:
                # Fallback: use Python's built-in csv module
                rows = []
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    reader = csv.reader(f)
                    for i, row in enumerate(reader):
                        if i > 50:
                            break
                        rows.append(" | ".join(row))
                summary = f"CSV File (first 50 rows):\n" + "\n".join(rows)
                logger.info(f"CSV extracted with fallback: {len(rows)} rows")
                return summary
        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
            raise ValueError(f"Could not extract CSV content: {e}")

    def _extract_xlsx(self, file_path: str) -> str:
        try:
            if HAS_PANDAS:
                xl = pd.ExcelFile(file_path)
                parts = [f"Excel file with sheets: {xl.sheet_names}"]
                for sheet in xl.sheet_names[:5]:
                    df = xl.parse(sheet, nrows=100)
                    parts.append(f"\n[Sheet: {sheet}]")
                    parts.append(df.head(20).to_string(index=False))
                return "\n".join(parts)
            else:
                # Fallback using openpyxl
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                parts = [f"Excel file with sheets: {wb.sheetnames}"]
                for sheet_name in wb.sheetnames[:3]:
                    ws = wb[sheet_name]
                    parts.append(f"\n[Sheet: {sheet_name}]")
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        if i > 20:
                            break
                        parts.append(" | ".join(str(c) if c is not None else "" for c in row))
                return "\n".join(parts)
        except Exception as e:
            logger.error(f"XLSX extraction error: {e}")
            raise ValueError(f"Could not extract XLSX content: {e}")

    def _extract_text(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            logger.info(f"TXT extracted: {len(text)} chars")
            return text
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise ValueError(f"Could not read text file: {e}")

    def chunk_text(self, text: str, chunk_size: int = 3000, overlap: int = 200) -> list:
        if len(text) <= chunk_size:
            return [text]
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start = end - overlap
        return chunks
